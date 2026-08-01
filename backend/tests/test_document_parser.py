from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
import pytest

from app.services.document_parser import (
    DocumentParseError,
    UnsupportedFileTypeError,
    extract_plain_text,
    sanitize_filename,
)


def test_extracts_utf8_txt_and_normalizes_whitespace() -> None:
    text = extract_plain_text("质量说明.txt", "第一段。\r\n\r\n\r\n第二段。".encode())
    assert text == "第一段。\n\n第二段。"


def test_extracts_gb18030_markdown() -> None:
    text = extract_plain_text("说明.md", "# 标题\n质量内容".encode("gb18030"))
    assert text == "# 标题\n质量内容"


def test_extracts_docx_paragraphs_and_tables() -> None:
    document = DocxDocument()
    document.add_paragraph("段落内容")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "缺陷"
    table.cell(0, 1).text = "偏移"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_plain_text("质量记录.docx", buffer.getvalue())

    assert "段落内容" in text
    assert "缺陷\t偏移" in text


def test_extracts_pdf_text() -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    resources = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 72 200 Td (QualityPilot PDF text) Tj ET")
    page[NameObject("/Resources")] = resources
    page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)

    text = extract_plain_text("manual.pdf", buffer.getvalue())

    assert "QualityPilot PDF text" in text


def test_rejects_unsupported_extension() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        extract_plain_text("data.xlsx", b"not an xlsx")


def test_rejects_binary_text_file() -> None:
    with pytest.raises(DocumentParseError, match="二进制"):
        extract_plain_text("data.txt", b"quality\x00pilot")


def test_sanitizes_path_components() -> None:
    assert sanitize_filename("../../documents/manual.txt") == "manual.txt"
    assert sanitize_filename(r"C:\\documents\\manual.txt") == "manual.txt"
