from io import BytesIO
from pathlib import PurePosixPath
import re

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class DocumentParseError(ValueError):
    """Raised when an uploaded document cannot produce usable plain text."""


class UnsupportedFileTypeError(DocumentParseError):
    """Raised when a filename has an unsupported extension."""


def sanitize_filename(filename: str | None) -> str:
    normalized = (filename or "").replace("\\", "/")
    safe_name = PurePosixPath(normalized).name.strip()
    if not safe_name or safe_name in {".", ".."}:
        raise DocumentParseError("文件名无效")
    if len(safe_name) > 255:
        raise DocumentParseError("文件名不能超过 255 个字符")
    return safe_name


def get_supported_extension(filename: str) -> str:
    suffix = PurePosixPath(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(f"不支持该文件类型，当前支持：{supported}")
    return suffix


def extract_plain_text(filename: str, content: bytes) -> str:
    extension = get_supported_extension(filename)
    if not content:
        raise DocumentParseError("文件内容为空")

    if extension in {".txt", ".md"}:
        text = _extract_text_file(content)
    elif extension == ".pdf":
        text = _extract_pdf(content)
    else:
        text = _extract_docx(content)

    normalized = normalize_text(text)
    if not normalized:
        raise DocumentParseError("文件中没有可提取的文本")
    return normalized


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized = "\n".join(lines).strip()
    return re.sub(r"\n{3,}", "\n\n", normalized)


def _extract_text_file(content: bytes) -> str:
    if b"\x00" in content:
        raise DocumentParseError("文本文件包含二进制内容")
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("文本文件必须使用 UTF-8 或 GB18030 编码")


def _extract_pdf(content: bytes) -> str:
    if not content.startswith(b"%PDF-"):
        raise DocumentParseError("PDF 文件签名无效")
    try:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            raise DocumentParseError("暂不支持加密 PDF")
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("PDF 解析失败") from exc
    return "\n\n".join(page for page in pages if page)


def _extract_docx(content: bytes) -> str:
    if not content.startswith(b"PK"):
        raise DocumentParseError("DOCX 文件签名无效")
    try:
        document = DocxDocument(BytesIO(content))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append("\t".join(cells))
    except Exception as exc:
        raise DocumentParseError("DOCX 解析失败") from exc
    return "\n\n".join(block for block in blocks if block)
